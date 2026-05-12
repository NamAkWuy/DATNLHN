"""
Đọc các file Markdown báo cáo (Chương 2, Chương 3-4...) → chèn vào cuối
ĐATN LÊ HOÀI NAM.docx, giữ style hiện có (Normal + bold cho heading,
Table Grid cho bảng).

Mặc định: chèn lần lượt Chương 2 (mục biểu đồ tuần tự / hoạt động) rồi
Chương 3 + 4. Có thể truyền danh sách file qua tham số dòng lệnh.

Chạy: python append_to_docx.py
       python append_to_docx.py BaoCao_Chuong_2_BieuDoTuanTu.md
"""
import re
import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

ROOT = Path(__file__).parent
DEFAULT_MD_FILES = [
    ROOT / "BaoCao_Chuong_2_BieuDoTuanTu.md",
    ROOT / "BaoCao_Chuong_3_4.md",
]
DOCX_FILE = ROOT / "ĐATN LÊ HOÀI NAM.docx"


def add_inline_runs(paragraph, text, italic=False):
    """Tách **bold**, *italic*, `code` rồi thêm run với định dạng tương ứng.
    italic match: dấu sao bao quanh, ngay sau dấu sao đầu KHÔNG phải khoảng trắng/sao
    (để không bắt nhầm các path kiểu admin/*).
    """
    pattern = re.compile(r"(\*\*.+?\*\*|\*[^\s*][^*]*?\*|`.+?`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run.text = part[1:-1]
            run.italic = True
        else:
            run.text = part
        if italic:
            run.italic = True


def add_chapter_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def add_h4(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(18)
    add_inline_runs(p, text, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run("• ")
    add_inline_runs(p, text)
    return p


def add_numbered(doc, num, text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.left_indent = Pt(18)
    p.add_run(f"{num}. ")
    add_inline_runs(p, text)
    return p


def add_quote(doc, text):
    """Blockquote → italic, indent."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(28)
    add_inline_runs(p, text, italic=True)


def add_image_placeholder(doc, caption):
    """Ô chèn ảnh: ghi rõ caption italic căn giữa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ Vị trí chèn ảnh — cậu chụp màn hình rồi paste vào đây ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(cap, caption, italic=True)


def add_mermaid_placeholder(doc, code, hint):
    """Block mermaid: ghi note rõ ràng để user export PNG paste vào."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"[ Sơ đồ Mermaid — {hint}. "
        "Mở file BaoCao_Chuong_3_4.md → copy code Mermaid → paste vào "
        "https://mermaid.live → Export PNG → chèn ảnh vào đây ]"
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_code_block(doc, code, lang=""):
    """Code block: font Consolas, nền xám nhạt qua paragraph spacing."""
    for line in code.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)


def parse_table_block(lines):
    """Đầu vào: list các dòng | a | b | (header, separator, body...).
    Trả về (header, body_rows)."""
    rows = []
    for ln in lines:
        ln = ln.strip()
        if not ln.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", ln):  # separator row
            continue
        cells = [c.strip() for c in ln.split("|")[1:-1]]
        rows.append(cells)
    if not rows:
        return [], []
    return rows[0], rows[1:]


def add_table(doc, header, body):
    n_cols = len(header)
    n_rows = 1 + len(body)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header
    for c_idx, cell_text in enumerate(header):
        cell = table.rows[0].cells[c_idx]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(cell_text)
        run.bold = True

    # Body
    for r_idx, row in enumerate(body, start=1):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= n_cols:
                break
            cell = table.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            # bỏ ký hiệu ✅ nếu user không có font hỗ trợ — giữ nguyên thử
            add_inline_runs(p, cell_text)

    # Empty paragraph after table
    doc.add_paragraph()


def parse_and_append(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ----- Bỏ qua dòng trống -----
        if not stripped:
            i += 1
            continue

        # ----- Horizontal rule -----
        if re.match(r"^-{3,}$", stripped):
            doc.add_paragraph()  # tạo khoảng trắng
            i += 1
            continue

        # ----- Heading -----
        if stripped.startswith("# CHƯƠNG") or stripped.startswith("# CHƯƠNG"):
            add_chapter_heading(doc, stripped[2:])
            i += 1
            continue
        # H1 cho file phụ trợ (vd. mục con của Chương 2)
        if stripped.startswith("# "):
            add_h2(doc, stripped[2:])
            i += 1
            continue
        if stripped.startswith("## "):
            add_h2(doc, stripped[3:])
            i += 1
            continue
        if stripped.startswith("### "):
            add_h3(doc, stripped[4:])
            i += 1
            continue
        if stripped.startswith("#### "):
            add_h4(doc, stripped[5:])
            i += 1
            continue

        # ----- Code block (```mermaid hoặc ```bash...) -----
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            code = "\n".join(code_lines)
            if lang == "mermaid":
                # Tìm hint từ caption ở dòng tiếp theo (Hình X.Y)
                hint_match = re.search(r"Hình \d+\.\d+", code) or "sơ đồ"
                first_word = code.split("\n")[0].strip() if code_lines else "diagram"
                # tìm tên sơ đồ trong dòng đầu code (flowchart, erDiagram, sequenceDiagram)
                if "erDiagram" in code:
                    hint = "ERD cơ sở dữ liệu"
                elif "sequenceDiagram" in code:
                    hint = "biểu đồ tuần tự (sequence)"
                elif "flowchart" in code:
                    hint = "biểu đồ luồng (flowchart)"
                else:
                    hint = "sơ đồ kỹ thuật"
                add_mermaid_placeholder(doc, code, hint)
            else:
                add_code_block(doc, code, lang)
            continue

        # ----- Table -----
        if stripped.startswith("|"):
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            header, body = parse_table_block(tbl_lines)
            if header:
                add_table(doc, header, body)
            continue

        # ----- Image placeholder line -----
        if "📷" in line or stripped.startswith("📷"):
            # dòng kiểu: 📷 *Ảnh chèn:* **Hình 3.5** — *Trang đăng nhập*
            # bỏ ký hiệu camera, lấy text còn lại
            clean = stripped.replace("📷", "").strip()
            # bỏ "*Ảnh chèn:*" và những dấu * thừa, lấy caption
            clean = re.sub(r"\*Ảnh chèn:\*", "", clean).strip()
            add_image_placeholder(doc, clean)
            i += 1
            continue

        # ----- Blockquote -----
        if stripped.startswith(">"):
            qtext = stripped.lstrip("> ").strip()
            # gom cả các dòng tiếp theo cũng bắt đầu bằng >
            i += 1
            while i < n and lines[i].strip().startswith(">"):
                qtext += " " + lines[i].strip().lstrip("> ").strip()
                i += 1
            add_quote(doc, qtext)
            continue

        # ----- Bullet list -----
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            add_bullet(doc, text)
            i += 1
            continue

        # ----- Numbered list -----
        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            add_numbered(doc, m.group(1), m.group(2))
            i += 1
            continue

        # ----- Paragraph thường -----
        # gom các dòng kế tiếp không phải heading/list/bảng
        body_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith(("#", "|", ">", "- ", "* ", "```", "📷"))
                or re.match(r"^\d+\.\s+", nxt)
                or re.match(r"^-{3,}$", nxt)
            ):
                break
            body_lines.append(nxt)
            i += 1
        add_body(doc, " ".join(body_lines))


def main():
    # CLI: nhận danh sách file md, mặc định dùng DEFAULT_MD_FILES
    if len(sys.argv) > 1:
        md_files = [ROOT / arg for arg in sys.argv[1:]]
    else:
        md_files = DEFAULT_MD_FILES

    for f in md_files:
        if not f.exists():
            print(f"❌ Không tìm thấy {f}")
            sys.exit(1)
    if not DOCX_FILE.exists():
        print(f"❌ Không tìm thấy {DOCX_FILE}")
        sys.exit(1)

    doc = Document(str(DOCX_FILE))
    before = len(doc.paragraphs)

    # Xóa các paragraph rỗng cuối file để nối liền mạch
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)

    from docx.enum.text import WD_BREAK
    for idx, md_file in enumerate(md_files):
        # Page break trước mỗi chương
        last = doc.add_paragraph()
        last.add_run().add_break(WD_BREAK.PAGE)

        md_text = md_file.read_text(encoding="utf-8")
        parse_and_append(doc, md_text)
        print(f"➕ Đã thêm: {md_file.name}")

    after = len(doc.paragraphs)

    out = DOCX_FILE
    doc.save(str(out))
    print(f"✅ Tổng cộng thêm {after - before} đoạn vào {out.name}")
    print(f"   Tổng paragraphs: {before} → {after}")
    print(f"   File size: {out.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
